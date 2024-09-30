import csv
import io
from typing import cast

import docx
import pandas as pd
import pypdfium2
from unstructured.partition.email import partition_email
from unstructured.partition.epub import partition_epub
from unstructured.partition.msg import partition_msg
from unstructured.partition.ppt import partition_ppt
from unstructured.partition.pptx import partition_pptx

from core.file import File, FileTransferMethod, file_manager
from core.helper import ssrf_proxy
from core.variables import ArrayFileSegment
from core.variables.segments import FileSegment
from core.workflow.entities.node_entities import NodeRunResult
from core.workflow.nodes.base_node import BaseNode
from enums import NodeType
from models.workflow import WorkflowNodeExecutionStatus

from .exc import DocumentExtractorError, FileDownloadError, TextExtractionError, UnsupportedFileTypeError
from .models import DocumentExtractorNodeData


class DocumentExtractorNode(BaseNode):
    """
    Extracts text content from various file types.
    Supports plain text, PDF, and DOC/DOCX files.
    """

    _node_data_cls = DocumentExtractorNodeData
    _node_type = NodeType.DOCUMENT_EXTRACTOR

    def _run(self):
        node_data = cast(DocumentExtractorNodeData, self.node_data)
        variable_selector = node_data.variable_selector
        variable = self.graph_runtime_state.variable_pool.get(variable_selector)

        if variable is None:
            error_message = f"File variable not found for selector: {variable_selector}"
            return NodeRunResult(status=WorkflowNodeExecutionStatus.FAILED, error=error_message)
        if variable.value and not isinstance(variable, ArrayFileSegment | FileSegment):
            error_message = f"Variable {variable_selector} is not an ArrayFileSegment"
            return NodeRunResult(status=WorkflowNodeExecutionStatus.FAILED, error=error_message)

        value = variable.value
        inputs = {"variable_selector": variable_selector}
        process_data = {"documents": value if isinstance(value, list) else [value]}

        try:
            if isinstance(value, list):
                extracted_text_list = list(map(_extract_text_from_file, value))
                return NodeRunResult(
                    status=WorkflowNodeExecutionStatus.SUCCEEDED,
                    inputs=inputs,
                    process_data=process_data,
                    outputs={"text": extracted_text_list},
                )
            elif isinstance(value, File):
                extracted_text = _extract_text_from_file(value)
                return NodeRunResult(
                    status=WorkflowNodeExecutionStatus.SUCCEEDED,
                    inputs=inputs,
                    process_data=process_data,
                    outputs={"text": extracted_text},
                )
            else:
                raise DocumentExtractorError(f"Unsupported variable type: {type(value)}")
        except DocumentExtractorError as e:
            return NodeRunResult(
                status=WorkflowNodeExecutionStatus.FAILED,
                error=str(e),
                inputs=inputs,
                process_data=process_data,
            )


def _extract_text(*, file_content: bytes, mime_type: str) -> str:
    """Extract text from a file based on its MIME type."""
    if mime_type.startswith("text/plain") or mime_type in {"text/html", "text/htm", "text/markdown", "text/xml"}:
        return _extract_text_from_plain_text(file_content)
    elif mime_type == "application/pdf":
        return _extract_text_from_pdf(file_content)
    elif mime_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        return _extract_text_from_doc(file_content)
    elif mime_type == "text/csv":
        return _extract_text_from_csv(file_content)
    elif mime_type in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    }:
        return _extract_text_from_excel(file_content)
    elif mime_type == "application/vnd.ms-powerpoint":
        return _extract_text_from_ppt(file_content)
    elif mime_type == "application/vnd.openxmlformats-officedocument.presentationml.presentation":
        return _extract_text_from_pptx(file_content)
    elif mime_type == "application/epub+zip":
        return _extract_text_from_epub(file_content)
    elif mime_type == "message/rfc822":
        return _extract_text_from_eml(file_content)
    elif mime_type == "application/vnd.ms-outlook":
        return _extract_text_from_msg(file_content)
    else:
        raise UnsupportedFileTypeError(f"Unsupported MIME type: {mime_type}")


def _extract_text_from_plain_text(file_content: bytes) -> str:
    try:
        return file_content.decode("utf-8")
    except UnicodeDecodeError as e:
        raise TextExtractionError("Failed to decode plain text file") from e


def _extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_document = pypdfium2.PdfDocument(pdf_file, autoclose=True)
        text = ""
        for page in pdf_document:
            text_page = page.get_textpage()
            text += text_page.get_text_range()
            text_page.close()
            page.close()
        return text
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}") from e


def _extract_text_from_doc(file_content: bytes) -> str:
    try:
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from DOC/DOCX: {str(e)}") from e


def _download_file_content(file: File) -> bytes:
    """Download the content of a file based on its transfer method."""
    try:
        if file.transfer_method == FileTransferMethod.REMOTE_URL:
            if file.remote_url is None:
                raise FileDownloadError("Missing URL for remote file")
            response = ssrf_proxy.get(file.remote_url)
            response.raise_for_status()
            return response.content
        elif file.transfer_method == FileTransferMethod.LOCAL_FILE:
            if file.related_id is None:
                raise FileDownloadError("Missing file ID for local file")
            return file_manager.download(upload_file_id=file.related_id, tenant_id=file.tenant_id)
        else:
            raise ValueError(f"Unsupported transfer method: {file.transfer_method}")
    except Exception as e:
        raise FileDownloadError(f"Error downloading file: {str(e)}") from e


def _extract_text_from_file(file: File):
    if file.mime_type is None:
        raise UnsupportedFileTypeError("Unable to determine file type: MIME type is missing")
    file_content = _download_file_content(file)
    extracted_text = _extract_text(file_content=file_content, mime_type=file.mime_type)
    return extracted_text


def _extract_text_from_csv(file_content: bytes) -> str:
    try:
        csv_file = io.StringIO(file_content.decode("utf-8"))
        csv_reader = csv.reader(csv_file)
        rows = list(csv_reader)

        if not rows:
            return ""

        # Create markdown table
        markdown_table = "| " + " | ".join(rows[0]) + " |\n"
        markdown_table += "| " + " | ".join(["---"] * len(rows[0])) + " |\n"
        for row in rows[1:]:
            markdown_table += "| " + " | ".join(row) + " |\n"

        return markdown_table.strip()
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from CSV: {str(e)}") from e


def _extract_text_from_excel(file_content: bytes) -> str:
    """Extract text from an Excel file using pandas."""

    try:
        df = pd.read_excel(io.BytesIO(file_content))

        # Drop rows where all elements are NaN
        df.dropna(how="all", inplace=True)

        # Convert DataFrame to markdown table
        markdown_table = df.to_markdown(index=False)
        return markdown_table
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from Excel file: {str(e)}") from e


def _extract_text_from_ppt(file_content: bytes) -> str:
    try:
        with io.BytesIO(file_content) as file:
            elements = partition_ppt(file=file)
        return "\n".join([getattr(element, "text", "") for element in elements])
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from PPT: {str(e)}") from e


def _extract_text_from_pptx(file_content: bytes) -> str:
    try:
        with io.BytesIO(file_content) as file:
            elements = partition_pptx(file=file)
        return "\n".join([getattr(element, "text", "") for element in elements])
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from PPTX: {str(e)}") from e


def _extract_text_from_epub(file_content: bytes) -> str:
    try:
        with io.BytesIO(file_content) as file:
            elements = partition_epub(file=file)
        return "\n".join([str(element) for element in elements])
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from EPUB: {str(e)}") from e


def _extract_text_from_eml(file_content: bytes) -> str:
    try:
        with io.BytesIO(file_content) as file:
            elements = partition_email(file=file)
        return "\n".join([str(element) for element in elements])
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from EML: {str(e)}") from e


def _extract_text_from_msg(file_content: bytes) -> str:
    try:
        with io.BytesIO(file_content) as file:
            elements = partition_msg(file=file)
        return "\n".join([str(element) for element in elements])
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from MSG: {str(e)}") from e